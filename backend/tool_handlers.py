"""
Tool Handlers for Anthropic Tool Use

Each handler receives:
- input_data: Dict with parameters from Claude's tool call
- context: ToolExecutionContext with session info (atividade_id, aluno_id, etc.)

Each handler returns:
- ToolResult with content string and is_error flag
"""

from typing import Dict, Any, Optional, List
import json
import re

from tools import ToolResult, ToolExecutionContext


# =============================================================================
# EXECUTE PYTHON CODE HANDLER
# =============================================================================

async def handle_execute_python_code(
    input_data: Dict[str, Any],
    context: Optional[ToolExecutionContext] = None
) -> ToolResult:
    """
    Handler for execute_python_code tool.

    Executes Python code in a sandboxed environment using the existing
    code_executor infrastructure (Docker or E2B).
    """
    from code_executor import (
        code_executor,
        ExecutionStatus,
        detect_libraries_from_code,
        detect_output_files_from_code
    )

    code = input_data.get("code", "")
    requested_files = input_data.get("output_files", [])
    description = input_data.get("description", "Code execution")

    if not code.strip():
        return ToolResult(
            tool_use_id="",
            content="Error: No code provided to execute.",
            is_error=True
        )

    try:
        # Detect libraries from code
        libraries = detect_libraries_from_code(code)

        # Detect output files from code (in addition to requested ones)
        detected_files = detect_output_files_from_code(code)
        all_output_files = list(set(requested_files + detected_files))

        # Execute the code
        result = await code_executor.execute(
            code=code,
            libraries=libraries,
            output_files=all_output_files if all_output_files else None
        )

        if result.is_success:
            # Format success response
            response_parts = []

            if description:
                response_parts.append(f"**Task:** {description}")

            response_parts.append("**Status:** Code executed successfully")

            if result.stdout and result.stdout.strip():
                response_parts.append(f"**Output:**\n```\n{result.stdout[:5000]}\n```")

            if result.files_generated:
                response_parts.append("**Files generated:**")
                files_info = []
                for f in result.files_generated:
                    size_kb = f.size_bytes / 1024
                    files_info.append({
                        "filename": f.filename,
                        "type": f.mime_type,
                        "size_kb": round(size_kb, 2),
                        "content_base64": f.content_base64[:200] + "..." if len(f.content_base64) > 200 else f.content_base64
                    })
                response_parts.append(f"```json\n{json.dumps(files_info, indent=2)}\n```")

            if result.plots_generated:
                response_parts.append(f"**Plots generated:** {len(result.plots_generated)} chart(s)")

            response_parts.append(f"**Execution time:** {result.execution_time_ms:.0f}ms")

            return ToolResult(
                tool_use_id="",
                content="\n\n".join(response_parts),
                is_error=False,
                files_generated=[f.to_dict() for f in result.files_generated]
            )

        elif result.status == ExecutionStatus.SECURITY_VIOLATION:
            return ToolResult(
                tool_use_id="",
                content=f"**Security Error:** The code was blocked due to security violations.\n\nDetails: {result.error_message}",
                is_error=True
            )

        elif result.status == ExecutionStatus.TIMEOUT:
            return ToolResult(
                tool_use_id="",
                content=f"**Timeout Error:** Code execution exceeded the time limit.\n\nPartial output:\n```\n{result.stdout[:1000]}\n```",
                is_error=True
            )

        else:
            # General error
            error_msg = result.stderr or result.error_message or "Unknown error"
            return ToolResult(
                tool_use_id="",
                content=f"**Execution Error:**\n```\n{error_msg[:3000]}\n```",
                is_error=True
            )

    except Exception as e:
        return ToolResult(
            tool_use_id="",
            content=f"**Internal Error:** Failed to execute code: {str(e)}",
            is_error=True
        )


# =============================================================================
# GET DOCUMENT CONTENT HANDLER
# =============================================================================

async def handle_get_document_content(
    input_data: Dict[str, Any],
    context: Optional[ToolExecutionContext] = None
) -> ToolResult:
    """
    Handler for get_document_content tool.

    Retrieves document content from storage.
    """
    from storage_v2 import storage_v2 as storage

    document_id = input_data.get("document_id")
    document_type = input_data.get("document_type")
    atividade_id = input_data.get("atividade_id") or (context.atividade_id if context else None)
    aluno_id = input_data.get("aluno_id") or (context.aluno_id if context else None)

    try:
        if document_id:
            # Get specific document by ID
            doc = storage.obter_documento(document_id)
            if not doc:
                return ToolResult(
                    tool_use_id="",
                    content=f"Document with ID '{document_id}' not found.",
                    is_error=True
                )

            # Read document content
            content = _read_document_content(doc.caminho_arquivo)
            return ToolResult(
                tool_use_id="",
                content=f"**Document:** {doc.nome_arquivo}\n**Type:** {doc.tipo.value}\n\n**Content:**\n{content[:10000]}",
                is_error=False
            )

        elif document_type and atividade_id:
            # Search by type and activity
            docs = storage.listar_documentos(atividade_id, aluno_id)
            matching = [d for d in docs if d.tipo.value == document_type]

            if not matching:
                return ToolResult(
                    tool_use_id="",
                    content=f"No documents of type '{document_type}' found for this activity.",
                    is_error=False
                )

            # Return first matching document
            doc = matching[0]
            content = _read_document_content(doc.caminho_arquivo)
            return ToolResult(
                tool_use_id="",
                content=f"**Document:** {doc.nome_arquivo}\n**Type:** {doc.tipo.value}\n\n**Content:**\n{content[:10000]}",
                is_error=False
            )

        else:
            return ToolResult(
                tool_use_id="",
                content="Please provide either document_id or (document_type + atividade_id)",
                is_error=True
            )

    except Exception as e:
        return ToolResult(
            tool_use_id="",
            content=f"Error retrieving document: {str(e)}",
            is_error=True
        )


def _read_document_content(file_path: str) -> str:
    """Read content from a document file"""
    from pathlib import Path
    import fitz  # PyMuPDF

    path = Path(file_path)

    if not path.exists():
        return "[File not found]"

    ext = path.suffix.lower()

    try:
        if ext in ['.txt', '.md', '.json', '.csv']:
            return path.read_text(encoding='utf-8')

        elif ext == '.pdf':
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text

        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])

        else:
            return f"[Cannot read file type: {ext}]"

    except Exception as e:
        return f"[Error reading file: {str(e)}]"


# =============================================================================
# GET STUDENT INFO HANDLER
# =============================================================================

async def handle_get_student_info(
    input_data: Dict[str, Any],
    context: Optional[ToolExecutionContext] = None
) -> ToolResult:
    """
    Handler for get_student_info tool.

    Retrieves student information from storage.
    """
    from storage_v2 import storage_v2 as storage

    aluno_id = input_data.get("aluno_id") or (context.aluno_id if context else None)
    include_grades = input_data.get("include_grades", True)

    if not aluno_id:
        return ToolResult(
            tool_use_id="",
            content="Please provide aluno_id (student ID)",
            is_error=True
        )

    try:
        aluno = storage.obter_aluno(aluno_id)

        if not aluno:
            return ToolResult(
                tool_use_id="",
                content=f"Student with ID '{aluno_id}' not found.",
                is_error=True
            )

        # Build response
        info = {
            "id": aluno.id,
            "nome": aluno.nome,
            "email": aluno.email,
            "turma_id": aluno.turma_id
        }

        response = f"**Student Information**\n"
        response += f"- **Name:** {aluno.nome}\n"
        response += f"- **Email:** {aluno.email or 'N/A'}\n"
        response += f"- **Class ID:** {aluno.turma_id}\n"

        if include_grades:
            # Get documents for this student (corrections, analyses)
            try:
                docs = storage.listar_documentos_por_aluno(aluno_id)
                corrections = [d for d in docs if 'correcao' in d.tipo.value.lower()]

                if corrections:
                    response += f"\n**Recent Activity:** {len(corrections)} correction(s) on record"
            except:
                pass

        return ToolResult(
            tool_use_id="",
            content=response,
            is_error=False
        )

    except Exception as e:
        return ToolResult(
            tool_use_id="",
            content=f"Error retrieving student info: {str(e)}",
            is_error=True
        )


# =============================================================================
# SEARCH DOCUMENTS HANDLER
# =============================================================================

async def handle_search_documents(
    input_data: Dict[str, Any],
    context: Optional[ToolExecutionContext] = None
) -> ToolResult:
    """
    Handler for search_documents tool.

    Searches documents by text content.
    """
    from storage_v2 import storage_v2 as storage

    query = input_data.get("query", "")
    document_type = input_data.get("document_type")
    limit = min(input_data.get("limit", 5), 20)

    if not query.strip():
        return ToolResult(
            tool_use_id="",
            content="Please provide a search query.",
            is_error=True
        )

    try:
        # Get all documents and filter
        all_docs = storage.listar_todos_documentos()

        if document_type and document_type != "all":
            all_docs = [d for d in all_docs if d.tipo.value == document_type]

        # Simple text search in filenames and content
        results = []
        query_lower = query.lower()

        for doc in all_docs[:100]:  # Limit search scope
            if query_lower in doc.nome_arquivo.lower():
                results.append(doc)
                continue

            # Try to read content for deeper search
            try:
                content = _read_document_content(doc.caminho_arquivo)
                if query_lower in content.lower():
                    results.append(doc)
            except:
                pass

            if len(results) >= limit:
                break

        if not results:
            return ToolResult(
                tool_use_id="",
                content=f"No documents found matching '{query}'",
                is_error=False
            )

        response = f"**Search Results for:** '{query}'\n\n"
        for i, doc in enumerate(results[:limit], 1):
            response += f"{i}. **{doc.nome_arquivo}** (Type: {doc.tipo.value})\n"
            response += f"   ID: {doc.id}\n"

        return ToolResult(
            tool_use_id="",
            content=response,
            is_error=False
        )

    except Exception as e:
        return ToolResult(
            tool_use_id="",
            content=f"Error searching documents: {str(e)}",
            is_error=True
        )


# =============================================================================
# SAVE CORRECTION HANDLER
# =============================================================================

async def handle_save_correction(
    input_data: Dict[str, Any],
    context: Optional[ToolExecutionContext] = None
) -> ToolResult:
    """
    Handler for save_correction tool.

    Saves grading/correction results to storage.
    """
    from storage_v2 import storage_v2 as storage
    from models import TipoDocumento
    import json as json_module
    from datetime import datetime

    questao_id = input_data.get("questao_id")
    aluno_id = input_data.get("aluno_id") or (context.aluno_id if context else None)
    nota = input_data.get("nota")
    nota_maxima = input_data.get("nota_maxima")
    feedback = input_data.get("feedback", "")
    erros = input_data.get("erros", [])

    # Validation
    if not all([questao_id, aluno_id, nota is not None, nota_maxima is not None]):
        return ToolResult(
            tool_use_id="",
            content="Missing required fields: questao_id, aluno_id, nota, nota_maxima",
            is_error=True
        )

    if nota < 0 or nota > nota_maxima:
        return ToolResult(
            tool_use_id="",
            content=f"Invalid score: {nota} (must be between 0 and {nota_maxima})",
            is_error=True
        )

    try:
        # Create correction document
        correction_data = {
            "questao_id": questao_id,
            "aluno_id": aluno_id,
            "nota": nota,
            "nota_maxima": nota_maxima,
            "percentual": round((nota / nota_maxima) * 100, 1),
            "feedback": feedback,
            "erros": erros,
            "data_correcao": datetime.now().isoformat(),
            "corrigido_por": "claude_tool"
        }

        # Save to storage
        atividade_id = context.atividade_id if context else None

        # Generate a unique filename
        filename = f"correcao_{questao_id}_{aluno_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        # Create temporary file and save
        import tempfile
        import os

        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, filename)

        with open(temp_path, 'w', encoding='utf-8') as f:
            json_module.dump(correction_data, f, ensure_ascii=False, indent=2)

        # Register in storage if atividade_id is available
        if atividade_id:
            doc = storage.criar_documento(
                nome_arquivo=filename,
                tipo=TipoDocumento.CORRECAO,
                atividade_id=atividade_id,
                aluno_id=aluno_id,
                caminho_arquivo=temp_path
            )

            return ToolResult(
                tool_use_id="",
                content=f"**Correction Saved**\n\n- Question: {questao_id}\n- Student: {aluno_id}\n- Score: {nota}/{nota_maxima} ({correction_data['percentual']}%)\n- Document ID: {doc.id}",
                is_error=False
            )
        else:
            return ToolResult(
                tool_use_id="",
                content=f"**Correction Recorded** (not saved to activity - no atividade_id)\n\n- Question: {questao_id}\n- Student: {aluno_id}\n- Score: {nota}/{nota_maxima} ({correction_data['percentual']}%)",
                is_error=False
            )

    except Exception as e:
        return ToolResult(
            tool_use_id="",
            content=f"Error saving correction: {str(e)}",
            is_error=True
        )


# =============================================================================
# CREATE DOCUMENT HANDLER
# =============================================================================

async def handle_create_document(
    input_data: Dict[str, Any],
    context: Optional[ToolExecutionContext] = None
) -> ToolResult:
    """
    Handler for create_document tool.

    Creates one or more documents and saves them to storage.
    Supports batch document creation for pipelines (e.g., creating reports for all students).
    """
    from storage_v2 import storage_v2 as storage
    from models import TipoDocumento
    from pathlib import Path
    from datetime import datetime
    import os
    import tempfile

    documents = input_data.get("documents", [])
    aluno_id = input_data.get("aluno_id") or (context.aluno_id if context else None)
    atividade_id = input_data.get("atividade_id") or (context.atividade_id if context else None)
    turma_id = input_data.get("turma_id")

    if not documents:
        return ToolResult(
            tool_use_id="",
            content="Error: No documents provided. Please specify at least one document to create.",
            is_error=True
        )

    # Map document types to TipoDocumento enum
    type_mapping = {
        "report": TipoDocumento.ANALISE,
        "feedback": TipoDocumento.CORRECAO,
        "analysis": TipoDocumento.ANALISE,
        "summary": TipoDocumento.ANALISE,
        "other": TipoDocumento.OUTROS,
        "exam": TipoDocumento.PROVA,
        "answer_key": TipoDocumento.GABARITO,
    }

    created_docs = []
    errors = []

    for idx, doc_data in enumerate(documents):
        try:
            filename = doc_data.get("filename", f"document_{idx+1}.txt")
            content = doc_data.get("content", "")
            doc_type_str = doc_data.get("document_type", "other").lower()
            description = doc_data.get("description", "")

            # Get TipoDocumento
            tipo = type_mapping.get(doc_type_str, TipoDocumento.OUTROS)

            # Determine file extension
            ext = Path(filename).suffix.lower()
            if not ext:
                ext = ".txt"
                filename = f"{filename}.txt"

            # Create the file content based on extension
            temp_dir = tempfile.gettempdir()
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_filename = f"{Path(filename).stem}_{timestamp}{ext}"
            temp_path = os.path.join(temp_dir, unique_filename)

            # Handle different file types
            if ext in ['.txt', '.md', '.csv', '.json']:
                # Plain text files
                with open(temp_path, 'w', encoding='utf-8') as f:
                    f.write(content)

            elif ext == '.pdf':
                # Create PDF using reportlab
                try:
                    from reportlab.lib.pagesizes import letter, A4
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.units import inch

                    doc = SimpleDocTemplate(temp_path, pagesize=A4)
                    styles = getSampleStyleSheet()
                    story = []

                    # Add title if description provided
                    if description:
                        title_style = ParagraphStyle(
                            'CustomTitle',
                            parent=styles['Heading1'],
                            fontSize=16,
                            spaceAfter=20
                        )
                        story.append(Paragraph(description, title_style))
                        story.append(Spacer(1, 12))

                    # Add content paragraphs
                    body_style = styles['Normal']
                    for paragraph in content.split('\n\n'):
                        if paragraph.strip():
                            story.append(Paragraph(paragraph.replace('\n', '<br/>'), body_style))
                            story.append(Spacer(1, 12))

                    doc.build(story)
                except ImportError:
                    # Fallback: save as text with .pdf.txt extension
                    temp_path = temp_path.replace('.pdf', '.pdf.txt')
                    unique_filename = unique_filename.replace('.pdf', '.pdf.txt')
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(f"[PDF content - reportlab not installed]\n\n{content}")

            elif ext == '.docx':
                # Create Word document
                try:
                    from docx import Document as DocxDocument
                    from docx.shared import Pt

                    doc = DocxDocument()

                    # Add title if description provided
                    if description:
                        doc.add_heading(description, level=0)

                    # Add content paragraphs
                    for paragraph in content.split('\n\n'):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph)

                    doc.save(temp_path)
                except ImportError:
                    # Fallback: save as text with .docx.txt extension
                    temp_path = temp_path.replace('.docx', '.docx.txt')
                    unique_filename = unique_filename.replace('.docx', '.docx.txt')
                    with open(temp_path, 'w', encoding='utf-8') as f:
                        f.write(f"[DOCX content - python-docx not installed]\n\n{content}")

            else:
                # Default: save as text
                with open(temp_path, 'w', encoding='utf-8') as f:
                    f.write(content)

            # Register document in storage if we have context
            doc_info = {
                "filename": unique_filename,
                "original_name": filename,
                "type": tipo.value,
                "description": description,
                "path": temp_path
            }

            if atividade_id:
                try:
                    saved_doc = storage.criar_documento(
                        nome_arquivo=unique_filename,
                        tipo=tipo,
                        atividade_id=atividade_id,
                        aluno_id=aluno_id,
                        caminho_arquivo=temp_path
                    )
                    doc_info["id"] = saved_doc.id
                    doc_info["saved_to_storage"] = True
                except Exception as e:
                    doc_info["saved_to_storage"] = False
                    doc_info["storage_error"] = str(e)
            else:
                doc_info["saved_to_storage"] = False
                doc_info["note"] = "No atividade_id provided - document created but not registered"

            created_docs.append(doc_info)

        except Exception as e:
            errors.append({
                "filename": doc_data.get("filename", f"document_{idx+1}"),
                "error": str(e)
            })

    # Build response
    response_parts = [f"**Documents Created: {len(created_docs)}/{len(documents)}**\n"]

    if created_docs:
        response_parts.append("**Successfully Created:**")
        for doc in created_docs:
            status = "✓ Saved to storage" if doc.get("saved_to_storage") else "⚠ File created (not in storage)"
            response_parts.append(f"- **{doc['original_name']}** ({doc['type']}) - {status}")
            if doc.get("id"):
                response_parts.append(f"  - Document ID: {doc['id']}")

    if errors:
        response_parts.append("\n**Errors:**")
        for err in errors:
            response_parts.append(f"- {err['filename']}: {err['error']}")

    # Include files_generated for the response
    files_generated = [
        {
            "filename": doc["filename"],
            "path": doc["path"],
            "type": doc["type"]
        }
        for doc in created_docs
    ]

    return ToolResult(
        tool_use_id="",
        content="\n".join(response_parts),
        is_error=len(errors) > 0 and len(created_docs) == 0,
        files_generated=files_generated
    )


# =============================================================================
# HANDLER REGISTRY
# =============================================================================

TOOL_HANDLERS = {
    "execute_python_code": handle_execute_python_code,
    "get_document_content": handle_get_document_content,
    "get_student_info": handle_get_student_info,
    "search_documents": handle_search_documents,
    "save_correction": handle_save_correction,
    "create_document": handle_create_document,
}


def get_handler(tool_name: str):
    """Get handler function for a tool"""
    return TOOL_HANDLERS.get(tool_name)
