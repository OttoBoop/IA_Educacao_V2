"""
Abstração de Providers de IA - Permite trocar facilmente entre diferentes modelos/APIs

Cada provider implementa a mesma interface, permitindo:
1. Trocar a IA usada em cada etapa do pipeline
2. Comparar resultados entre diferentes IAs
3. Rastrear qual IA gerou cada output
"""

from abc import ABC, abstractmethod
from typing import Optional, Dict, Any, List
from dataclasses import dataclass, field
from datetime import datetime
import json
import os
from pathlib import Path

# Importar gerenciador de chaves criptografadas (importação adiada para evitar circular)
_api_key_manager = None

def _get_api_key_manager():
    """Obtém instância do gerenciador de chaves (lazy loading para evitar import circular)"""
    global _api_key_manager
    if _api_key_manager is None:
        from chat_service import api_key_manager, ProviderType
        _api_key_manager = api_key_manager
    return _api_key_manager

def _get_provider_type():
    """Obtém enum ProviderType (lazy loading)"""
    from chat_service import ProviderType
    return ProviderType




def _testing_mode() -> bool:
    """Detecta se esta rodando em contexto de testes."""
    return os.getenv("PROVA_AI_TESTING", "").lower() in ("1", "true", "yes")


def _local_llm_disabled() -> bool:
    """Permite desabilitar provider local (ex: Ollama) via env var."""
    return os.getenv("PROVA_AI_DISABLE_LOCAL_LLM", "").lower() in ("1", "true", "yes")
@dataclass
class AIResponse:
    """Resposta padronizada de qualquer provider de IA"""
    content: str
    provider: str
    model: str
    tokens_used: int
    latency_ms: float
    input_tokens: int = 0
    output_tokens: int = 0
    timestamp: datetime = field(default_factory=datetime.now)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "provider": self.provider,
            "model": self.model,
            "tokens_used": self.tokens_used,
            "input_tokens": self.input_tokens,
            "output_tokens": self.output_tokens,
            "latency_ms": self.latency_ms,
            "timestamp": self.timestamp.isoformat(),
            "metadata": self.metadata
        }


def _format_httpx_error(exc) -> str:
    response = exc.response
    request = exc.request
    body = response.text.strip()
    if len(body) > 2000:
        body = f"{body[:2000]}... (truncado)"
    request_id = response.headers.get("x-request-id") or response.headers.get("request-id")
    request_id_info = f" request_id={request_id}" if request_id else ""
    return (
        f"Erro HTTP {response.status_code} em {request.method} {request.url}.{request_id_info} "
        f"Resposta: {body}"
    )


class AIProvider(ABC):
    """Interface base para todos os providers de IA"""

    def __init__(self, api_key: str, model: str):
        self.api_key = api_key
        self.model = model
        self.name = self.__class__.__name__

    @abstractmethod
    async def complete(self,
                       prompt: str,
                       system_prompt: Optional[str] = None,
                       temperature: float = 0.7,
                       max_tokens: int = 4096,
                       reasoning_effort: Optional[str] = None) -> AIResponse:
        """Gera uma completion dado um prompt

        Args:
            prompt: O prompt do usuário
            system_prompt: Prompt de sistema opcional
            temperature: Controle de criatividade (0-2). Ignorado para modelos de raciocínio
            max_tokens: Número máximo de tokens na resposta
            reasoning_effort: Para modelos de raciocínio (o1, o3, o4-mini): 'low', 'medium' ou 'high'
        """
        pass
    
    @abstractmethod
    async def analyze_document(self, 
                               file_path: str,
                               instruction: str) -> AIResponse:
        """Analisa um documento (PDF, DOCX, imagem)"""
        pass
    
    def get_identifier(self) -> str:
        """Retorna identificador único para rastreamento"""
        return f"{self.name}_{self.model}"


class OpenAIProvider(AIProvider):
    """Provider para OpenAI (GPT-4, GPT-4o, GPT-5, o1, o3, o4-mini, etc.)"""

    # Modelos de raciocínio que não suportam temperature
    # Deprecated: 'o1', 'o1-mini', 'o1-pro' (removidos pois estão deprecated)
    REASONING_MODELS = {'o3', 'o3-mini', 'o3-pro', 'o4-mini', 'gpt-5', 'gpt-5-mini', 'gpt-5-nano', 'gpt-5.1', 'gpt-5.2'}

    def __init__(self, api_key: str, model: str = "gpt-4o"):
        super().__init__(api_key, model)
        self.base_url = "https://api.openai.com/v1"

    def _is_reasoning_model(self) -> bool:
        """Verifica se o modelo é um modelo de raciocínio"""
        return self.model in self.REASONING_MODELS

    async def complete(self,
                       prompt: str,
                       system_prompt: Optional[str] = None,
                       temperature: float = 0.7,
                       max_tokens: int = 4096,
                       reasoning_effort: Optional[str] = None) -> AIResponse:
        import httpx
        import time

        start = time.time()

        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})

        # Construir payload base
        payload = {
            "model": self.model,
            "messages": messages,
        }

        # Modelos de raciocínio: usar reasoning_effort em vez de temperature
        if self._is_reasoning_model():
            # Reasoning models usam max_completion_tokens, não max_tokens
            payload["max_completion_tokens"] = max_tokens
            # Definir reasoning_effort - usar "minimal" por padrão para evitar respostas vazias
            # Ref: https://community.openai.com/t/gpt-5-mini-api-unstable-and-slow-repeated-timeout-and-empty-responses/1355041
            # Modelos GPT-5 podem consumir todos tokens em raciocínio se effort for alto demais
            effort = reasoning_effort if reasoning_effort in ['minimal', 'low', 'medium', 'high'] else 'minimal'
            payload["reasoning_effort"] = effort
        else:
            # Modelos normais: usar temperature e max_tokens
            payload["temperature"] = temperature
            payload["max_tokens"] = max_tokens

        import asyncio
        async with httpx.AsyncClient() as client:
            max_retries = 3
            backoff = 1.0
            data = None
            content = ""

            for attempt in range(max_retries):
                try:
                    response = await client.post(
                        f"{self.base_url}/chat/completions",
                        headers={
                            "Authorization": f"Bearer {self.api_key}",
                            "Content-Type": "application/json"
                        },
                        json=payload,
                        timeout=120.0
                    )
                    response.raise_for_status()
                    data = response.json()

                    # Extrair conteúdo - pode ser None para modelos de raciocínio
                    # Ref: https://github.com/openai/openai-python/issues/2546
                    content = data["choices"][0]["message"].get("content") or ""

                    # Se resposta vazia em modelo de raciocínio, retry com effort maior
                    # Ref: https://community.openai.com/t/gpt-5-mini-api-unstable-and-slow-repeated-timeout-and-empty-responses/1355041
                    if content or not self._is_reasoning_model():
                        break  # Sucesso

                    # Retry com reasoning_effort mais alto
                    if attempt < max_retries - 1:
                        await asyncio.sleep(backoff)
                        backoff *= 2
                        payload["reasoning_effort"] = "low" if attempt == 0 else "medium"
                        continue

                except httpx.HTTPStatusError as exc:
                    status = exc.response.status_code
                    if status in (429, 500, 502, 503, 504) and attempt < max_retries - 1:
                        await asyncio.sleep(backoff)
                        backoff *= 2
                        continue
                    raise RuntimeError(_format_httpx_error(exc)) from exc

        latency = (time.time() - start) * 1000

        return AIResponse(
            content=content,
            provider="openai",
            model=self.model,
            tokens_used=data["usage"]["total_tokens"],
            input_tokens=data["usage"].get("prompt_tokens", 0),
            output_tokens=data["usage"].get("completion_tokens", 0),
            latency_ms=latency,
            metadata={
                "finish_reason": data["choices"][0]["finish_reason"],
                "is_reasoning_model": self._is_reasoning_model(),
                "reasoning_tokens": data["usage"].get("completion_tokens_details", {}).get("reasoning_tokens", 0)
            }
        )
    
    async def analyze_document(self, 
                               file_path: str,
                               instruction: str) -> AIResponse:
        import httpx
        import base64
        import time
        
        start = time.time()
        
        # Determinar tipo de arquivo
        ext = Path(file_path).suffix.lower()
        
        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            # Imagem - usar vision
            with open(file_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode()
            
            mime_type = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }[ext]
            
            messages = [{
                "role": "user",
                "content": [
                    {"type": "text", "text": instruction},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{image_data}"
                        }
                    }
                ]
            }]
        else:
            # Texto - extrair conteúdo primeiro
            content = await self._extract_text(file_path)
            messages = [{
                "role": "user",
                "content": f"{instruction}\n\n---\nConteúdo do documento:\n{content}"
            }]
        
        async with httpx.AsyncClient() as client:
            try:
                response = await client.post(
                    f"{self.base_url}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": self.model,
                        "messages": messages,
                        "max_tokens": 4096
                    },
                    timeout=120.0
                )
                response.raise_for_status()
                data = response.json()
            except httpx.HTTPStatusError as exc:
                raise RuntimeError(_format_httpx_error(exc)) from exc
        
        latency = (time.time() - start) * 1000
        
        return AIResponse(
            content=data["choices"][0]["message"]["content"],
            provider="openai",
            model=self.model,
            tokens_used=data["usage"]["total_tokens"],
            input_tokens=data["usage"].get("prompt_tokens", 0),
            output_tokens=data["usage"].get("completion_tokens", 0),
            latency_ms=latency,
            metadata={
                "file_analyzed": file_path,
                "finish_reason": data["choices"][0]["finish_reason"]
            }
        )
    
    async def _extract_text(self, file_path: str) -> str:
        """Extrai texto de diferentes formatos"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.pdf':
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError(f"Formato não suportado: {ext}")


class AnthropicProvider(AIProvider):
    """Provider para Anthropic (Claude)"""

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-20250514"):
        super().__init__(api_key, model)
        self.base_url = "https://api.anthropic.com/v1"

    async def complete(self,
                       prompt: str,
                       system_prompt: Optional[str] = None,
                       temperature: float = 0.7,
                       max_tokens: int = 4096,
                       reasoning_effort: Optional[str] = None,
                       tools: Optional[List[Dict[str, Any]]] = None,
                       tool_choice: Optional[Dict[str, Any]] = None) -> AIResponse:
        import httpx
        import time

        start = time.time()

        payload = {
            "model": self.model,
            "max_tokens": max_tokens,
            "temperature": temperature,  # Claude suporta temperature 0-1
            "messages": [{"role": "user", "content": prompt}]
        }
        if system_prompt:
            payload["system"] = system_prompt
        if tools:
            payload["tools"] = tools
        if tool_choice:
            payload["tool_choice"] = tool_choice

        async with httpx.AsyncClient() as client:
            try:
                response = await client.post(
                    f"{self.base_url}/messages",
                    headers={
                        "x-api-key": self.api_key,
                        "anthropic-version": "2023-06-01",
                        "Content-Type": "application/json"
                    },
                    json=payload,
                    timeout=120.0
                )
                response.raise_for_status()
                data = response.json()
            except httpx.HTTPStatusError as exc:
                raise RuntimeError(_format_httpx_error(exc)) from exc

        latency = (time.time() - start) * 1000

        # Extract text content (may not exist if only tool_use blocks)
        text_content = ""
        for block in data.get("content", []):
            if block.get("type") == "text":
                text_content += block.get("text", "")

        return AIResponse(
            content=text_content,
            provider="anthropic",
            model=self.model,
            tokens_used=data["usage"]["input_tokens"] + data["usage"]["output_tokens"],
            input_tokens=data["usage"].get("input_tokens", 0),
            output_tokens=data["usage"].get("output_tokens", 0),
            latency_ms=latency,
            metadata={
                "stop_reason": data["stop_reason"],
                "content_blocks": data.get("content", [])  # Include full content for tool handling
            }
        )
    
    async def analyze_document(self, 
                               file_path: str,
                               instruction: str) -> AIResponse:
        import httpx
        import base64
        import time
        
        start = time.time()
        ext = Path(file_path).suffix.lower()
        
        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            with open(file_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode()
            
            mime_type = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }[ext]
            
            content = [
                {"type": "text", "text": instruction},
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": mime_type,
                        "data": image_data
                    }
                }
            ]
        elif ext == '.pdf':
            with open(file_path, 'rb') as f:
                pdf_data = base64.b64encode(f.read()).decode()
            
            content = [
                {"type": "text", "text": instruction},
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_data
                    }
                }
            ]
        else:
            # Extrair texto para outros formatos
            text_content = await self._extract_text(file_path)
            content = [{"type": "text", "text": f"{instruction}\n\n---\n{text_content}"}]
        
        async with httpx.AsyncClient() as client:
            try:
                response = await client.post(
                    f"{self.base_url}/messages",
                    headers={
                        "x-api-key": self.api_key,
                        "anthropic-version": "2023-06-01",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": self.model,
                        "max_tokens": 4096,
                        "messages": [{"role": "user", "content": content}]
                    },
                    timeout=120.0
                )
                response.raise_for_status()
                data = response.json()
            except httpx.HTTPStatusError as exc:
                raise RuntimeError(_format_httpx_error(exc)) from exc
        
        latency = (time.time() - start) * 1000
        
        return AIResponse(
            content=data["content"][0]["text"],
            provider="anthropic",
            model=self.model,
            tokens_used=data["usage"]["input_tokens"] + data["usage"]["output_tokens"],
            input_tokens=data["usage"].get("input_tokens", 0),
            output_tokens=data["usage"].get("output_tokens", 0),
            latency_ms=latency,
            metadata={
                "file_analyzed": file_path,
                "stop_reason": data["stop_reason"]
            }
        )
    
    async def _extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError(f"Formato não suportado para extração de texto: {ext}")


class GeminiProvider(AIProvider):
    """Provider para Google Gemini (Gemini 2.5 Pro, Flash, etc.)"""

    def __init__(self, api_key: str, model: str = "gemini-2.5-flash"):
        super().__init__(api_key, model)
        self.base_url = "https://generativelanguage.googleapis.com/v1beta"

    async def complete(self,
                       prompt: str,
                       system_prompt: Optional[str] = None,
                       temperature: float = 0.7,
                       max_tokens: int = 4096,
                       reasoning_effort: Optional[str] = None) -> AIResponse:
        import httpx
        import time

        start = time.time()

        # Construir conteúdo com system instruction se fornecido
        contents = []
        
        # Adicionar mensagem do usuário
        contents.append({
            "role": "user",
            "parts": [{"text": prompt}]
        })

        # Construir payload
        payload = {
            "contents": contents,
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens
            }
        }

        # System instruction separada (Gemini API v1beta style)
        if system_prompt:
            payload["systemInstruction"] = {
                "parts": [{"text": system_prompt}]
            }

        async with httpx.AsyncClient() as client:
            try:
                response = await client.post(
                    f"{self.base_url}/models/{self.model}:generateContent",
                    headers={
                        "x-goog-api-key": self.api_key,
                        "Content-Type": "application/json"
                    },
                    json=payload,
                    timeout=120.0
                )
                response.raise_for_status()
                data = response.json()
            except httpx.HTTPStatusError as exc:
                raise RuntimeError(_format_httpx_error(exc)) from exc

        latency = (time.time() - start) * 1000

        # Extrair texto da resposta
        content_text = ""
        if "candidates" in data and len(data["candidates"]) > 0:
            candidate = data["candidates"][0]
            if "content" in candidate and "parts" in candidate["content"]:
                parts = candidate["content"]["parts"]
                content_text = "".join(p.get("text", "") for p in parts)

        # Extrair uso de tokens
        usage = data.get("usageMetadata", {})
        input_tokens = usage.get("promptTokenCount", 0)
        output_tokens = usage.get("candidatesTokenCount", 0)

        return AIResponse(
            content=content_text,
            provider="google",
            model=self.model,
            tokens_used=input_tokens + output_tokens,
            input_tokens=input_tokens,
            output_tokens=output_tokens,
            latency_ms=latency,
            metadata={
                "finish_reason": data.get("candidates", [{}])[0].get("finishReason", "unknown")
            }
        )

    async def analyze_document(self,
                               file_path: str,
                               instruction: str) -> AIResponse:
        import httpx
        import base64
        import time

        start = time.time()

        # Determinar tipo de arquivo
        ext = Path(file_path).suffix.lower()

        contents = []
        parts = []

        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
            # Imagem - usar inline_data
            with open(file_path, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode()

            mime_type = {
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif',
                '.webp': 'image/webp'
            }.get(ext, 'image/png')

            parts.append({
                "inline_data": {
                    "mime_type": mime_type,
                    "data": image_data
                }
            })
            parts.append({"text": instruction})

        elif ext == '.pdf':
            # PDF - Gemini suporta PDFs nativamente
            with open(file_path, 'rb') as f:
                pdf_data = base64.b64encode(f.read()).decode()

            parts.append({
                "inline_data": {
                    "mime_type": "application/pdf",
                    "data": pdf_data
                }
            })
            parts.append({"text": instruction})

        else:
            # Outros formatos - extrair texto
            text_content = await self._extract_text(file_path)
            parts.append({"text": f"{instruction}\n\n---\n{text_content}"})

        contents.append({
            "role": "user",
            "parts": parts
        })

        payload = {
            "contents": contents,
            "generationConfig": {
                "maxOutputTokens": 4096
            }
        }

        async with httpx.AsyncClient() as client:
            try:
                response = await client.post(
                    f"{self.base_url}/models/{self.model}:generateContent",
                    headers={
                        "x-goog-api-key": self.api_key,
                        "Content-Type": "application/json"
                    },
                    json=payload,
                    timeout=120.0
                )
                response.raise_for_status()
                data = response.json()
            except httpx.HTTPStatusError as exc:
                raise RuntimeError(_format_httpx_error(exc)) from exc

        latency = (time.time() - start) * 1000

        # Extrair texto da resposta
        content_text = ""
        if "candidates" in data and len(data["candidates"]) > 0:
            candidate = data["candidates"][0]
            if "content" in candidate and "parts" in candidate["content"]:
                parts = candidate["content"]["parts"]
                content_text = "".join(p.get("text", "") for p in parts)

        usage = data.get("usageMetadata", {})
        input_tokens = usage.get("promptTokenCount", 0)
        output_tokens = usage.get("candidatesTokenCount", 0)

        return AIResponse(
            content=content_text,
            provider="google",
            model=self.model,
            tokens_used=input_tokens + output_tokens,
            input_tokens=input_tokens,
            output_tokens=output_tokens,
            latency_ms=latency,
            metadata={
                "file_analyzed": file_path,
                "finish_reason": data.get("candidates", [{}])[0].get("finishReason", "unknown")
            }
        )

    async def _extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError(f"Formato não suportado para extração de texto: {ext}")


class LocalLLMProvider(AIProvider):
    """Provider para LLMs locais via Ollama ou similar"""

    def __init__(self, base_url: str = "http://localhost:11434", model: str = "llama3"):
        super().__init__("", model)
        self.base_url = base_url

    async def complete(self,
                       prompt: str,
                       system_prompt: Optional[str] = None,
                       temperature: float = 0.7,
                       max_tokens: int = 4096,
                       reasoning_effort: Optional[str] = None) -> AIResponse:
        import httpx
        import time
        
        start = time.time()
        
        full_prompt = prompt
        if system_prompt:
            full_prompt = f"{system_prompt}\n\n{prompt}"
        
        async with httpx.AsyncClient() as client:
            response = await client.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": self.model,
                    "prompt": full_prompt,
                    "stream": False,
                    "options": {
                        "temperature": temperature,
                        "num_predict": max_tokens
                    }
                },
                timeout=300.0
            )
            response.raise_for_status()
            data = response.json()
        
        latency = (time.time() - start) * 1000
        
        return AIResponse(
            content=data["response"],
            provider="ollama",
            model=self.model,
            tokens_used=data.get("eval_count", 0) + data.get("prompt_eval_count", 0),
            input_tokens=data.get("prompt_eval_count", 0),
            output_tokens=data.get("eval_count", 0),
            latency_ms=latency,
            metadata={"done": data.get("done", True)}
        )
    
    async def analyze_document(self, 
                               file_path: str,
                               instruction: str) -> AIResponse:
        # Para modelos locais, extrair texto primeiro
        text = await self._extract_text(file_path)
        return await self.complete(f"{instruction}\n\n---\n{text}")
    
    async def _extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif ext == '.pdf':
            import fitz
            doc = fitz.open(file_path)
            return "\n".join([page.get_text() for page in doc])
        elif ext == '.docx':
            from docx import Document
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError(f"Formato não suportado: {ext}")


class AIProviderRegistry:
    """Registro central de providers - permite trocar IAs facilmente"""

    def __init__(self, config_path: str = "./data/providers.json"):
        self.providers: Dict[str, AIProvider] = {}
        self.provider_configs: Dict[str, Dict[str, Any]] = {}  # Configs para persistência
        self.default_provider: Optional[str] = None
        self.provider_health: Dict[str, str] = {}  # name -> "ok" | "error: ..."
        self.config_path = Path(config_path)
        self.config_path.parent.mkdir(parents=True, exist_ok=True)
        self._load_from_file()

    def _load_from_file(self):
        """Carrega providers salvos do arquivo"""
        if not self.config_path.exists():
            return

        try:
            with open(self.config_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            self.default_provider = data.get("default_provider")

            for config in data.get("providers", []):
                try:
                    provider_type = (config.get("provider_type") or "").lower()
                    if _local_llm_disabled() and provider_type in ("ollama", "localllmprovider"):
                        continue
                    self._create_provider_from_config(config)
                except Exception as e:
                    print(f"Erro ao carregar provider {config.get('name')}: {e}")

            if self.default_provider and self.default_provider not in self.providers:
                self.default_provider = None

            print(f"[OK] Providers carregados: {list(self.providers.keys())}")
        except Exception as e:
            print(f"Erro ao carregar providers do arquivo: {e}")

    def _save_to_file(self):
        """Salva providers para o arquivo"""
        try:
            data = {
                "default_provider": self.default_provider,
                "providers": list(self.provider_configs.values())
            }
            with open(self.config_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Erro ao salvar providers: {e}")

    def _create_provider_from_config(self, config: Dict[str, Any]):
        """Cria um provider a partir de uma configuração salva"""
        name = config["name"]
        provider_type = config["provider_type"].lower()
        model = config["model"]
        base_url = config.get("base_url")

        # Buscar API key do gerenciador de chaves criptografadas
        api_key = self._get_api_key_for_provider(provider_type)

        if provider_type == "openai" or provider_type == "openaiprovider":
            provider = OpenAIProvider(api_key=api_key, model=model)
        elif provider_type == "anthropic" or provider_type == "anthropicprovider":
            provider = AnthropicProvider(api_key=api_key, model=model)
        elif provider_type in ("google", "gemini", "geminiprovider"):
            provider = GeminiProvider(api_key=api_key, model=model)
        elif provider_type == "ollama" or provider_type == "localllmprovider":
            provider = LocalLLMProvider(
                base_url=base_url or "http://localhost:11434",
                model=model
            )
        else:
            raise ValueError(f"Tipo de provider desconhecido: {provider_type}")

        self.providers[name] = provider
        # Não armazenar api_key na config
        config_sem_key = {k: v for k, v in config.items() if k != "api_key"}
        self.provider_configs[name] = config_sem_key

    def _get_api_key_for_provider(self, provider_type: str) -> str:
        """Busca API key do gerenciador de chaves criptografadas"""
        try:
            ProviderType = _get_provider_type()
            key_manager = _get_api_key_manager()

            # Mapear tipo de provider para ProviderType enum
            type_mapping = {
                "openai": ProviderType.OPENAI,
                "openaiprovider": ProviderType.OPENAI,
                "anthropic": ProviderType.ANTHROPIC,
                "anthropicprovider": ProviderType.ANTHROPIC,
                "google": ProviderType.GOOGLE,
                "gemini": ProviderType.GOOGLE,
                "geminiprovider": ProviderType.GOOGLE,
            }

            provider_enum = type_mapping.get(provider_type.lower())
            if provider_enum:
                key_config = key_manager.get_por_empresa(provider_enum)
                if key_config:
                    return key_config.api_key

            # Fallback para variáveis de ambiente
            env_vars = {
                "openai": "OPENAI_API_KEY",
                "openaiprovider": "OPENAI_API_KEY",
                "anthropic": "ANTHROPIC_API_KEY",
                "anthropicprovider": "ANTHROPIC_API_KEY",
                "google": "GOOGLE_API_KEY",
                "gemini": "GOOGLE_API_KEY",
                "geminiprovider": "GOOGLE_API_KEY",
            }
            env_var = env_vars.get(provider_type.lower())
            if env_var:
                return os.getenv(env_var, "")

            return ""
        except Exception as e:
            print(f"[WARN] Erro ao buscar API key para {provider_type}: {e}")
            return ""

    def register(self, name: str, provider: AIProvider, set_default: bool = False,
                 api_key: str = None, base_url: str = None):
        """Registra um provider com um nome e persiste a configuração

        NOTA: O parâmetro api_key é ignorado. As chaves são gerenciadas pelo
        sistema centralizado em api_keys.json (criptografado).
        """
        self.providers[name] = provider

        # Salvar configuração para persistência (SEM api_key - gerenciado separadamente)
        self.provider_configs[name] = {
            "name": name,
            "provider_type": provider.name,
            "model": provider.model,
            "base_url": base_url or getattr(provider, 'base_url', None)
        }

        if set_default or self.default_provider is None:
            self.default_provider = name

        self._save_to_file()

    def unregister(self, name: str) -> bool:
        """Remove um provider"""
        if name in self.providers:
            del self.providers[name]
            del self.provider_configs[name]

            # Se era o default, escolhe outro
            if self.default_provider == name:
                self.default_provider = next(iter(self.providers.keys()), None)

            self._save_to_file()
            return True
        return False

    def set_default(self, name: str) -> bool:
        """Define o provider padrão"""
        if name in self.providers:
            self.default_provider = name
            self._save_to_file()
            return True
        return False

    def get(self, name: Optional[str] = None) -> AIProvider:
        """Obtém um provider pelo nome ou retorna o default"""
        if name is None:
            name = self.default_provider
        if name not in self.providers:
            raise ValueError(f"Provider '{name}' não registrado")
        return self.providers[name]

    def get_default(self) -> Optional[AIProvider]:
        """Retorna o provider padrao, ou None se nenhum estiver configurado"""
        if self.default_provider and self.default_provider in self.providers:
            provider = self.providers[self.default_provider]
            if _testing_mode():
                # Em testes, evita usar providers sem API key ou LLM local desativado.
                if isinstance(provider, LocalLLMProvider) and _local_llm_disabled():
                    return None
                if not isinstance(provider, LocalLLMProvider) and not getattr(provider, 'api_key', ''):
                    return None
            return provider
        return None

    def list_providers(self) -> List[str]:
        """Lista todos os providers registrados"""
        return list(self.providers.keys())

    def get_provider_info(self) -> List[Dict[str, str]]:
        """Retorna informações de todos os providers"""
        return [
            {
                "name": name,
                "provider_type": p.name,
                "model": p.model,
                "identifier": p.get_identifier()
            }
            for name, p in self.providers.items()
        ]


# Instância global do registry
ai_registry = AIProviderRegistry()


def setup_providers_from_env():
    """Configura providers a partir de variáveis de ambiente ou api_keys.json.

    Só adiciona providers se não houver nenhum carregado do arquivo.
    Isso permite que os providers salvos tenham prioridade.

    NOTA: As API keys são gerenciadas pelo sistema centralizado em api_keys.json
    (criptografado). Variáveis de ambiente são usadas apenas como fallback.
    """
    # Se já existem providers carregados do arquivo, não sobrescrever
    if ai_registry.providers:
        print(f"[INFO] Usando {len(ai_registry.providers)} provider(s) do arquivo de configuração")
        return

    print("[INFO] Nenhum provider salvo encontrado, configurando a partir do ambiente/keys...")

    # Tentar obter chaves do gerenciador criptografado
    try:
        ProviderType = _get_provider_type()
        key_manager = _get_api_key_manager()

        # OpenAI
        openai_config = key_manager.get_por_empresa(ProviderType.OPENAI)
        openai_key = openai_config.api_key if openai_config else os.getenv("OPENAI_API_KEY", "")

        if openai_key:
            ai_registry.register(
                "openai-gpt4o",
                OpenAIProvider(openai_key, "gpt-4o"),
            )
            ai_registry.register(
                "openai-gpt4o-mini",
                OpenAIProvider(openai_key, "gpt-4o-mini")
            )

        # Anthropic — haiku is the preferred default for the pipeline
        anthropic_config = key_manager.get_por_empresa(ProviderType.ANTHROPIC)
        anthropic_key = anthropic_config.api_key if anthropic_config else os.getenv("ANTHROPIC_API_KEY", "")

        if anthropic_key:
            ai_registry.register(
                "claude-sonnet",
                AnthropicProvider(anthropic_key, "claude-sonnet-4-20250514")
            )
            ai_registry.register(
                "claude-haiku",
                AnthropicProvider(anthropic_key, "claude-haiku-4-5-20251001"),
                set_default=True
            )

    except Exception as e:
        print(f"[WARN] Erro ao carregar chaves do gerenciador: {e}")

    # Ollama (local) - sempre disponível como fallback
    if not _local_llm_disabled() and "ollama-llama3" not in ai_registry.providers:
        ai_registry.register(
            "ollama-llama3",
            LocalLLMProvider(model="llama3"),
            base_url="http://localhost:11434"
        )

    # Health check: validate default provider can make API calls
    _validate_default_provider()


def _validate_default_provider():
    """Validate that the default provider can accept API calls.

    If the default provider fails, try to fall back to another provider
    and log a clear warning so the issue is visible in deployment logs.
    Health results are stored in ai_registry.provider_health.
    """
    import asyncio

    if os.getenv("PROVA_AI_TESTING") == "1":
        return  # Skip health check during tests

    default_name = ai_registry.default_provider
    if not default_name or default_name not in ai_registry.providers:
        print("[CRITICAL] No default AI provider configured!")
        ai_registry.provider_health["_status"] = "no_default"
        return

    async def _ping(prov):
        try:
            resp = await prov.complete(
                "Reply with exactly: OK",
                max_tokens=5,
                temperature=0.0,
            )
            return resp.content.strip() if resp else None
        except Exception as e:
            return f"error: {e}"

    def _run_ping(prov):
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                return None  # Can't run sync from async context
            return loop.run_until_complete(_ping(prov))
        except RuntimeError:
            loop = asyncio.new_event_loop()
            result = loop.run_until_complete(_ping(prov))
            loop.close()
            return result

    # Check all non-ollama providers
    for name, prov in ai_registry.providers.items():
        if name.startswith("ollama"):
            continue
        result = _run_ping(prov)
        if result is None:
            ai_registry.provider_health[name] = "deferred"
            print(f"[INFO] Provider '{name}' health check deferred (async context)")
        elif result and "OK" in str(result).upper()[:10]:
            ai_registry.provider_health[name] = "ok"
            print(f"[OK] Provider '{name}' ({prov.model}) is healthy")
        else:
            ai_registry.provider_health[name] = f"error: {result}"
            print(f"[WARN] Provider '{name}' ({prov.model}) FAILED: {result}")

    # Check if default is healthy
    default_health = ai_registry.provider_health.get(default_name, "unknown")
    if default_health == "ok":
        return

    # Default provider failed — try to fall back
    print(f"[WARN] Default provider '{default_name}' is unhealthy: {default_health}")
    print(f"[WARN] Attempting to find a working fallback provider...")

    for name in ai_registry.providers:
        if name == default_name or name.startswith("ollama"):
            continue
        if ai_registry.provider_health.get(name) == "ok":
            original_default = default_name
            ai_registry.set_default(name)
            print(f"[WARN] Fell back to '{name}' as default (was '{original_default}')")
            print(f"[WARN] FIX: Update the API key for '{original_default}' to restore it as default")
            return

    print(f"[CRITICAL] No working AI providers found! All providers failed health check.")
