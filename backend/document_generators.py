"""
NOVO CR - Document Generators v1.0

Módulo para geração de documentos em múltiplos formatos:
- PDF (via reportlab)
- CSV (via csv module)
- DOCX (via python-docx)

Cada função recebe dados estruturados (Dict, VisaoAluno, etc.)
e retorna bytes ou string pronto para salvar.
"""

import csv
import io
import json
from datetime import datetime
from typing import Dict, Any, List, Optional, Union
from dataclasses import dataclass
from enum import Enum


class OutputFormat(Enum):
    """Formatos de saída suportados"""
    JSON = "json"
    PDF = "pdf"
    CSV = "csv"
    DOCX = "docx"
    MD = "md"


# Mapeamento: tipo de documento do pipeline → formatos de saída
# Configura quais formatos gerar para cada tipo de resultado
STAGE_OUTPUT_FORMATS: Dict[str, List[OutputFormat]] = {
    "extracao_questoes": [OutputFormat.JSON],
    "extracao_respostas": [OutputFormat.JSON],
    "correcao": [OutputFormat.JSON, OutputFormat.PDF],
    "analise_habilidades": [OutputFormat.JSON, OutputFormat.PDF, OutputFormat.CSV],
    "relatorio_final": [OutputFormat.JSON, OutputFormat.PDF],
    "ranking": [OutputFormat.JSON, OutputFormat.CSV],
}


def get_output_formats(stage_type: str) -> List[OutputFormat]:
    """Retorna os formatos configurados para um tipo de documento"""
    return STAGE_OUTPUT_FORMATS.get(stage_type, [OutputFormat.JSON])


# ============================================================
# PDF GENERATION
# ============================================================

def generate_pdf(data: Dict[str, Any], title: str = "Documento", 
                 doc_type: str = "generic") -> bytes:
    """
    Gera PDF a partir de dados estruturados.
    
    Args:
        data: Dicionário com dados do documento
        title: Título do documento
        doc_type: Tipo de documento (correcao, relatorio_final, analise_habilidades)
    
    Returns:
        bytes: Conteúdo do PDF
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    except ImportError:
        # Fallback: retorna texto se reportlab não disponível
        return _generate_text_fallback(data, title).encode('utf-8')
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                           leftMargin=2*cm, rightMargin=2*cm,
                           topMargin=2*cm, bottomMargin=2*cm)
    
    styles = getSampleStyleSheet()
    
    # Estilos customizados
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#2c3e50')
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10,
        textColor=colors.HexColor('#34495e')
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        leading=14
    )
    
    story = []

    # Título
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))

    # Check for pipeline error — add prominent error section at the top
    if "_erro_pipeline" in data:
        erro = data["_erro_pipeline"]
        erro_style = ParagraphStyle(
            'ErroTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.red,
            alignment=TA_CENTER,
            spaceBefore=10,
            spaceAfter=15
        )
        story.append(Paragraph("ERRO DE PROCESSAMENTO", erro_style))

        erro_table_data = [
            ["Campo", "Valor"],
            ["Tipo", str(erro.get("tipo", "-"))],
            ["Mensagem", str(erro.get("mensagem", "-"))],
            ["Etapa", str(erro.get("etapa", "-"))],
            ["Severidade", str(erro.get("severidade", "-"))],
        ]
        erro_table = Table(erro_table_data, colWidths=[3*cm, 12*cm])
        erro_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.red),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fff0f0')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(erro_table)
        story.append(Spacer(1, 20))

    # Geração baseada no tipo
    if doc_type == "correcao":
        story.extend(_build_correcao_pdf(data, styles, heading_style, body_style))
    elif doc_type == "relatorio_final":
        story.extend(_build_relatorio_pdf(data, styles, heading_style, body_style))
    elif doc_type == "analise_habilidades":
        story.extend(_build_analise_pdf(data, styles, heading_style, body_style))
    else:
        story.extend(_build_generic_pdf(data, styles, heading_style, body_style))
    
    # Footer com timestamp
    story.append(Spacer(1, 30))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], 
                                   fontSize=8, textColor=colors.gray)
    story.append(Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _build_correcao_pdf(data: Dict, styles, heading_style, body_style) -> List:
    """Constrói conteúdo PDF para correção"""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    
    story = []
    
    # Resumo
    if "nota" in data:
        nota = data.get("nota", 0)
        nota_max = data.get("nota_maxima", 10)
        percentual = (nota / nota_max * 100) if nota_max > 0 else 0
        
        story.append(Paragraph("📊 Resumo", heading_style))
        summary_data = [
            ["Nota", f"{nota:.1f} / {nota_max:.1f}"],
            ["Percentual", f"{percentual:.1f}%"],
            ["Status", data.get("status", "N/A").upper()]
        ]
        table = Table(summary_data, colWidths=[150, 200])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(table)
        story.append(Spacer(1, 15))
    
    # Correções por questão
    if "correcoes" in data:
        story.append(Paragraph("📝 Questões", heading_style))
        for i, c in enumerate(data["correcoes"], 1):
            q_num = c.get("questao_numero", i)
            story.append(Paragraph(f"<b>Questão {q_num}</b>", body_style))
            story.append(Paragraph(f"Nota: {c.get('nota', 0):.1f} / {c.get('nota_maxima', 1):.1f}", body_style))
            if c.get("feedback"):
                story.append(Paragraph(f"<i>{c.get('feedback')}</i>", body_style))
            story.append(Spacer(1, 10))
    
    # Feedback geral
    if data.get("feedback"):
        story.append(Paragraph("💬 Feedback", heading_style))
        story.append(Paragraph(data["feedback"], body_style))
    
    return story


def _build_relatorio_pdf(data: Dict, styles, heading_style, body_style) -> List:
    """Constrói conteúdo PDF para relatório final"""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    
    story = []
    
    # Info do aluno
    if data.get("aluno_nome"):
        story.append(Paragraph(f"<b>Aluno:</b> {data.get('aluno_nome')}", body_style))
    if data.get("atividade_nome"):
        story.append(Paragraph(f"<b>Atividade:</b> {data.get('atividade_nome')}", body_style))
    story.append(Spacer(1, 15))
    
    # Nota final
    story.append(Paragraph("Resultado Final", heading_style))
    nota = data.get("nota_final", data.get("nota", 0)) or 0
    nota_max = data.get("nota_maxima", 10) or 10
    percentual = data.get("percentual") or ((nota / nota_max * 100) if nota_max > 0 else 0)
    
    summary_data = [
        ["Nota Final", f"{nota:.1f} / {nota_max:.1f}"],
        ["Aproveitamento", f"{percentual:.1f}%"],
    ]
    
    if data.get("total_questoes"):
        summary_data.append(["Total de Questões", str(data["total_questoes"])])
        summary_data.append(["Corretas", str(data.get("questoes_corretas", 0))])
        summary_data.append(["Parciais", str(data.get("questoes_parciais", 0))])
        summary_data.append(["Incorretas", str(data.get("questoes_incorretas", 0))])
    
    table = Table(summary_data, colWidths=[150, 200])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
        ('PADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(table)
    story.append(Spacer(1, 20))
    
    # Feedback geral
    if data.get("feedback_geral"):
        story.append(Paragraph("Comentarios", heading_style))
        story.append(Paragraph(data["feedback_geral"], body_style))

    # Recomendações
    if data.get("recomendacoes"):
        story.append(Paragraph("Recomendacoes", heading_style))
        for rec in data["recomendacoes"]:
            story.append(Paragraph(f"• {rec}", body_style))
    
    return story


def _build_analise_pdf(data: Dict, styles, heading_style, body_style) -> List:
    """Constrói conteúdo PDF para análise de habilidades"""
    from reportlab.platypus import Paragraph, Spacer
    
    story = []
    
    # Habilidades demonstradas
    habilidades = data.get("habilidades", {})
    
    if isinstance(habilidades, dict):
        dominadas = habilidades.get("dominadas", [])
        em_dev = habilidades.get("em_desenvolvimento", [])
        nao_dem = habilidades.get("nao_demonstradas", [])
    else:
        dominadas = data.get("habilidades_demonstradas", [])
        em_dev = []
        nao_dem = data.get("habilidades_faltantes", [])
    
    if dominadas:
        story.append(Paragraph("✅ Habilidades Demonstradas", heading_style))
        for h in dominadas:
            if isinstance(h, dict):
                story.append(Paragraph(f"• {h.get('nome', h)}", body_style))
            else:
                story.append(Paragraph(f"• {h}", body_style))
        story.append(Spacer(1, 10))
    
    if em_dev:
        story.append(Paragraph("🔄 Em Desenvolvimento", heading_style))
        for h in em_dev:
            if isinstance(h, dict):
                story.append(Paragraph(f"• {h.get('nome', h)}", body_style))
            else:
                story.append(Paragraph(f"• {h}", body_style))
        story.append(Spacer(1, 10))
    
    if nao_dem:
        story.append(Paragraph("⚠️ Precisam de Atenção", heading_style))
        for h in nao_dem:
            if isinstance(h, dict):
                story.append(Paragraph(f"• {h.get('nome', h)}", body_style))
            else:
                story.append(Paragraph(f"• {h}", body_style))
    
    return story


def _build_generic_pdf(data: Dict, styles, heading_style, body_style) -> List:
    """Constrói conteúdo PDF genérico para qualquer estrutura"""
    from reportlab.platypus import Paragraph, Spacer
    
    story = []
    
    for key, value in data.items():
        if key.startswith("_"):
            continue
            
        label = key.replace("_", " ").title()
        
        if isinstance(value, (list, dict)):
            story.append(Paragraph(f"<b>{label}:</b>", body_style))
            story.append(Paragraph(f"<pre>{json.dumps(value, indent=2, ensure_ascii=False)[:500]}</pre>", body_style))
        else:
            story.append(Paragraph(f"<b>{label}:</b> {value}", body_style))
        
        story.append(Spacer(1, 5))
    
    return story


def narrative_markdown_to_pdf(md_text: str, title: str = "Documento") -> bytes:
    """
    Convert narrative Markdown to a styled PDF via HTML intermediate.

    Used by the two-pass pipeline: Pass 2 generates rich Markdown narrative,
    this function converts it to a professional PDF for the professor.

    Args:
        md_text: Markdown text from the narrative prompt
        title: Document title (e.g., "Correção Narrativa — Maria Silva")

    Returns:
        bytes: PDF content
    """
    try:
        import markdown as md_lib
        from xhtml2pdf import pisa

        html_body = md_lib.markdown(md_text or "", extensions=['extra'])
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/>
<style>
  @page {{ size: A4; margin: 2cm; }}
  body {{ font-family: Helvetica, Arial, sans-serif; line-height: 1.6; color: #333; }}
  h1 {{ color: #2c3e50; text-align: center; margin-bottom: 1em; font-size: 20pt; }}
  h2 {{ color: #34495e; border-bottom: 1px solid #ddd; padding-bottom: 4px; margin-top: 1.2em; font-size: 14pt; }}
  h3 {{ color: #555; margin-top: 1em; font-size: 12pt; }}
  p {{ margin-bottom: 0.5em; font-size: 11pt; }}
  ul, ol {{ margin-left: 1.5em; font-size: 11pt; }}
  li {{ margin-bottom: 0.3em; }}
  strong {{ color: #2c3e50; }}
  blockquote {{ border-left: 3px solid #3498db; padding-left: 1em; color: #555; margin: 0.8em 0; }}
  .footer {{ font-size: 8pt; color: #999; margin-top: 2em; text-align: right; }}
</style>
</head><body>
  <h1>{title}</h1>
  {html_body}
  <p class="footer">Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
</body></html>"""

        buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            full_html.encode('utf-8'),
            dest=buffer,
            encoding='utf-8'
        )
        buffer.seek(0)
        return buffer.read()

    except ImportError:
        # Fallback: use reportlab with basic text rendering
        return _narrative_reportlab_fallback(md_text, title)


def _narrative_reportlab_fallback(md_text: str, title: str) -> bytes:
    """Fallback PDF generation using ReportLab when xhtml2pdf is not available."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('NarrTitle', parent=styles['Heading1'],
                                      fontSize=18, alignment=TA_CENTER,
                                      textColor=colors.HexColor('#2c3e50'))
        heading_style = ParagraphStyle('NarrHeading', parent=styles['Heading2'],
                                        fontSize=14, textColor=colors.HexColor('#34495e'))
        body_style = ParagraphStyle('NarrBody', parent=styles['Normal'],
                                     fontSize=11, leading=14, spaceAfter=6)

        story = [Paragraph(title, title_style), Spacer(1, 12)]

        for line in (md_text or "").split('\n'):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 6))
            elif stripped.startswith('## '):
                story.append(Paragraph(stripped[3:], heading_style))
            elif stripped.startswith('### '):
                story.append(Paragraph(stripped[4:], body_style))
            elif stripped.startswith('- ') or stripped.startswith('* '):
                story.append(Paragraph(f"\u2022 {stripped[2:]}", body_style))
            else:
                text = stripped.replace('**', '<b>').replace('**', '</b>')
                story.append(Paragraph(text, body_style))

        footer_style = ParagraphStyle('Footer', parent=styles['Normal'],
                                       fontSize=8, textColor=colors.gray)
        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    except ImportError:
        # Last resort: plain text as bytes
        text = f"{title}\n{'='*50}\n\n{md_text or ''}\n\nGerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        return text.encode('utf-8')


def generate_pipeline_pdf(
    tipo: str,
    json_data: Dict[str, Any],
    header: Optional[Dict[str, str]] = None,
) -> bytes:
    """Server-side PDF generation for pipeline artifacts (CORRECAO, ANALISE_HABILIDADES, RELATORIO_FINAL).

    Uses the exact field names from the pipeline JSON schema (`nota_final`,
    `questoes`, `feedback_geral`, etc.) and renders text in the labels the
    cross-check validator expects ("Nota final: X", "Questão N - Nota: M",
    "Feedback Geral: ..."). Replaces the model-generated PDF that the
    pipeline used to require via execute_python_code.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    header = header or {}
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "PipelineTitle", parent=styles["Heading1"], fontSize=18,
        spaceAfter=14, alignment=TA_CENTER, textColor=colors.HexColor("#2c3e50"),
    )
    heading_style = ParagraphStyle(
        "PipelineHeading", parent=styles["Heading2"], fontSize=13,
        spaceBefore=12, spaceAfter=8, textColor=colors.HexColor("#34495e"),
    )
    body_style = ParagraphStyle(
        "PipelineBody", parent=styles["Normal"], fontSize=11, leading=15, spaceAfter=6,
    )
    label_style = ParagraphStyle(
        "PipelineLabel", parent=styles["Normal"], fontSize=11, leading=15,
        spaceAfter=4, textColor=colors.HexColor("#2c3e50"),
    )

    story: List[Any] = []
    titles = {
        "correcao": "Correção da Atividade",
        "analise_habilidades": "Análise de Habilidades",
        "relatorio_final": "Relatório Final do Aluno",
    }
    story.append(Paragraph(titles.get(tipo, "Documento"), title_style))

    def _hdr(label: str, value: Optional[str]) -> None:
        if value:
            story.append(Paragraph(f"<b>{label}:</b> {value}", label_style))

    _hdr("Aluno", header.get("aluno_nome"))
    _hdr("Matéria", header.get("materia_nome"))
    _hdr("Atividade", header.get("atividade_nome"))
    _hdr("Data", header.get("data") or datetime.now().strftime("%d/%m/%Y"))
    story.append(Spacer(1, 10))

    if tipo == "correcao":
        story.extend(_build_pipeline_correcao(json_data, heading_style, body_style))
    elif tipo == "analise_habilidades":
        story.extend(_build_pipeline_analise(json_data, heading_style, body_style))
    elif tipo == "relatorio_final":
        story.extend(_build_pipeline_relatorio(json_data, heading_style, body_style))
    else:
        story.extend(_build_generic_pdf(json_data, styles, heading_style, body_style))

    story.append(Spacer(1, 18))
    footer_style = ParagraphStyle(
        "PipelineFooter", parent=styles["Normal"], fontSize=8, textColor=colors.gray,
    )
    story.append(Paragraph(
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} (auto-render server-side)",
        footer_style,
    ))
    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _fmt_num(v: Any) -> str:
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v) if v is not None else "—"
    return str(int(f)) if f.is_integer() else f"{f:.2f}".rstrip("0").rstrip(".")


def _escape(text: Any) -> str:
    if text is None:
        return ""
    s = str(text)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_pipeline_correcao(data: Dict, heading_style, body_style) -> List:
    from reportlab.platypus import Paragraph, Spacer
    story = []
    nota_final = data.get("nota_final")
    if nota_final is not None:
        story.append(Paragraph(f"<b>Nota final:</b> {_fmt_num(nota_final)}", body_style))
    total_a = data.get("total_acertos")
    total_e = data.get("total_erros")
    if total_a is not None or total_e is not None:
        story.append(Paragraph(
            f"Acertos: {_fmt_num(total_a or 0)} | Erros: {_fmt_num(total_e or 0)}",
            body_style,
        ))
    story.append(Spacer(1, 8))

    questoes = data.get("questoes") or []
    if questoes:
        story.append(Paragraph("Questões", heading_style))
        for q in questoes:
            if not isinstance(q, dict):
                continue
            numero = q.get("numero", "?")
            rc = str(q.get("resposta_correta") or "").strip().upper()
            uncorrectable = rc in ("", "MISSING_CONTENT", "N/A", "NULL", "NONE")
            if uncorrectable:
                header_line = f"<b>Questão {numero}</b> — Sem gabarito disponível"
            else:
                nota = q.get("nota")
                nota_max = q.get("nota_maxima")
                acerto = q.get("acerto")
                ac_label = "✓" if acerto else ("✗" if acerto is False else "?")
                header_line = (
                    f"<b>Questão {numero}</b> {ac_label} — Nota: {_fmt_num(nota)} / "
                    f"{_fmt_num(nota_max) if nota_max is not None else '10'}"
                )
            story.append(Paragraph(header_line, body_style))
            ra = q.get("resposta_aluno")
            if ra:
                story.append(Paragraph(f"<i>Resposta do aluno:</i> {_escape(ra)}", body_style))
            if not uncorrectable and q.get("resposta_correta"):
                story.append(Paragraph(
                    f"<i>Resposta correta:</i> {_escape(q.get('resposta_correta'))}", body_style,
                ))
            if q.get("feedback"):
                story.append(Paragraph(f"<i>Feedback:</i> {_escape(q.get('feedback'))}", body_style))
            story.append(Spacer(1, 6))

    feedback_geral = data.get("feedback_geral")
    if feedback_geral:
        story.append(Paragraph("Feedback Geral", heading_style))
        story.append(Paragraph(_escape(feedback_geral), body_style))

    avisos = data.get("_avisos_documento") or []
    if avisos:
        story.append(Spacer(1, 8))
        story.append(Paragraph("Avisos", heading_style))
        for a in avisos:
            if isinstance(a, dict):
                story.append(Paragraph(
                    f"• [{a.get('codigo','?')}] {_escape(a.get('explicacao',''))}", body_style,
                ))
    return story


def _build_pipeline_analise(data: Dict, heading_style, body_style) -> List:
    from reportlab.platypus import Paragraph, Spacer
    story = []
    habilidades = data.get("habilidades") or []
    if habilidades:
        story.append(Paragraph("Habilidades Avaliadas", heading_style))
        for h in habilidades:
            if not isinstance(h, dict):
                continue
            nome = h.get("nome", "—")
            nivel = h.get("nivel", "—")
            nota = h.get("nota")
            nota_label = f" — Nota: {_fmt_num(nota)}" if nota is not None else ""
            story.append(Paragraph(
                f"<b>{_escape(nome)}</b> — Nível: {_escape(nivel)}{nota_label}", body_style,
            ))
            evidencias = h.get("evidencias") or []
            for e in evidencias:
                story.append(Paragraph(f"  • {_escape(e)}", body_style))
            story.append(Spacer(1, 4))

    indicadores = data.get("indicadores") or {}
    if isinstance(indicadores, dict) and indicadores:
        story.append(Paragraph("Indicadores", heading_style))
        prof = indicadores.get("proficiencia_geral")
        if prof is not None:
            story.append(Paragraph(f"Proficiência geral: {_fmt_num(prof)}", body_style))
        for label_key, sec_label in (
            ("areas_destaque", "Áreas de destaque"),
            ("areas_atencao", "Áreas de atenção"),
        ):
            items = indicadores.get(label_key) or []
            if items:
                story.append(Paragraph(f"<b>{sec_label}:</b>", body_style))
                for it in items:
                    story.append(Paragraph(f"  • {_escape(it)}", body_style))

    recomendacoes = data.get("recomendacoes") or []
    if recomendacoes:
        story.append(Paragraph("Recomendações", heading_style))
        for r in recomendacoes:
            if isinstance(r, dict):
                line = f"[{_escape(r.get('prioridade','?'))}] {_escape(r.get('tipo',''))}: {_escape(r.get('descricao',''))}"
            else:
                line = f"• {_escape(r)}"
            story.append(Paragraph(line, body_style))
    return story


def _build_pipeline_relatorio(data: Dict, heading_style, body_style) -> List:
    from reportlab.platypus import Paragraph, Spacer
    story = []
    nota_final = data.get("nota_final")
    if nota_final is not None:
        story.append(Paragraph(f"<b>Nota final:</b> {_fmt_num(nota_final)}", body_style))
        story.append(Spacer(1, 6))

    resumo = data.get("resumo_geral")
    if resumo:
        story.append(Paragraph("Resumo Geral", heading_style))
        story.append(Paragraph(_escape(resumo), body_style))

    fortes = data.get("pontos_fortes") or []
    if fortes:
        story.append(Paragraph("Pontos Fortes", heading_style))
        for p in fortes:
            story.append(Paragraph(f"• {_escape(p)}", body_style))

    melhorias = data.get("areas_melhoria") or []
    if melhorias:
        story.append(Paragraph("Áreas para Melhoria", heading_style))
        for m in melhorias:
            story.append(Paragraph(f"• {_escape(m)}", body_style))

    recomendacoes = data.get("recomendacoes") or []
    if recomendacoes:
        story.append(Paragraph("Recomendações", heading_style))
        for r in recomendacoes:
            if isinstance(r, dict):
                line = f"[{_escape(r.get('prioridade','?'))}] {_escape(r.get('tipo',''))}: {_escape(r.get('descricao',''))}"
            else:
                line = f"• {_escape(r)}"
            story.append(Paragraph(line, body_style))

    detalhamento = data.get("detalhamento")
    if detalhamento:
        story.append(Paragraph("Detalhamento", heading_style))
        story.append(Paragraph(_escape(detalhamento), body_style))
    return story


def _generate_text_fallback(data: Dict, title: str) -> str:
    """Fallback para texto quando reportlab não está disponível"""
    lines = [f"{'='*50}", f" {title}", f"{'='*50}", ""]

    # Error section if pipeline error present
    if "_erro_pipeline" in data:
        erro = data["_erro_pipeline"]
        lines.append("!!! ERRO DE PROCESSAMENTO !!!")
        lines.append(f"Tipo: {erro.get('tipo', '-')}")
        lines.append(f"Mensagem: {erro.get('mensagem', '-')}")
        lines.append(f"Etapa: {erro.get('etapa', '-')}")
        lines.append(f"Severidade: {erro.get('severidade', '-')}")
        lines.append("")

    for key, value in data.items():
        if key.startswith("_"):
            continue
        label = key.replace("_", " ").title()
        if isinstance(value, (list, dict)):
            lines.append(f"{label}:")
            lines.append(json.dumps(value, indent=2, ensure_ascii=False))
        else:
            lines.append(f"{label}: {value}")
    
    lines.extend(["", f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}"])
    return "\n".join(lines)


# ============================================================
# CSV GENERATION
# ============================================================

def generate_csv(data: Union[List[Dict], Dict], 
                 headers: Optional[List[str]] = None,
                 doc_type: str = "generic") -> str:
    """
    Gera CSV a partir de dados estruturados.
    
    Args:
        data: Lista de dicionários ou dicionário único
        headers: Cabeçalhos opcionais (inferidos automaticamente)
        doc_type: Tipo de documento para formatação específica
    
    Returns:
        str: Conteúdo CSV
    """
    output = io.StringIO()
    
    # Normalizar para lista
    if isinstance(data, dict):
        # Caso especial: ranking ou lista dentro do dict
        if "ranking" in data:
            rows = data["ranking"]
        elif "correcoes" in data:
            rows = data["correcoes"]
        elif "habilidades" in data:
            rows = _flatten_habilidades(data["habilidades"])
        else:
            rows = [data]
    else:
        rows = data
    
    if not rows:
        return ""
    
    # Inferir headers
    if not headers:
        if rows and isinstance(rows[0], dict):
            headers = list(rows[0].keys())
        else:
            headers = ["value"]
            rows = [{"value": r} for r in rows]
    
    # Escrever CSV
    writer = csv.DictWriter(output, fieldnames=headers, extrasaction='ignore')
    writer.writeheader()
    
    for row in rows:
        if isinstance(row, dict):
            # Flatten nested values
            flat_row = {}
            for k, v in row.items():
                if isinstance(v, (list, dict)):
                    flat_row[k] = json.dumps(v, ensure_ascii=False)
                else:
                    flat_row[k] = v
            writer.writerow(flat_row)
        else:
            writer.writerow({"value": row})
    
    return output.getvalue()


def _flatten_habilidades(habilidades: Dict) -> List[Dict]:
    """Transforma habilidades em lista para CSV"""
    rows = []
    
    if isinstance(habilidades, dict):
        for status, lista in habilidades.items():
            if isinstance(lista, list):
                for h in lista:
                    if isinstance(h, dict):
                        rows.append({"status": status, **h})
                    else:
                        rows.append({"status": status, "habilidade": h})
    
    return rows


# ============================================================
# MARKDOWN GENERATION
# ============================================================

def generate_markdown(data: Dict[str, Any], title: str = "Documento",
                      doc_type: str = "generic") -> str:
    """
    Gera Markdown a partir de dados estruturados.
    
    Args:
        data: Dicionário com dados
        title: Título do documento
        doc_type: Tipo de documento
    
    Returns:
        str: Conteúdo Markdown
    """
    lines = [f"# {title}", ""]
    
    if doc_type == "correcao":
        lines.extend(_build_correcao_md(data))
    elif doc_type == "relatorio_final":
        lines.extend(_build_relatorio_md(data))
    elif doc_type == "analise_habilidades":
        lines.extend(_build_analise_md(data))
    else:
        lines.extend(_build_generic_md(data))
    
    lines.extend(["", "---", f"*Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}*"])
    return "\n".join(lines)


def _build_correcao_md(data: Dict) -> List[str]:
    """Constrói Markdown para correção"""
    lines = []
    
    if "nota" in data:
        nota = data.get("nota", 0)
        nota_max = data.get("nota_maxima", 10)
        lines.extend([
            "## 📊 Resumo",
            "",
            f"| Métrica | Valor |",
            f"|---------|-------|",
            f"| Nota | {nota:.1f} / {nota_max:.1f} |",
            f"| Status | {data.get('status', 'N/A').upper()} |",
            ""
        ])
    
    if "correcoes" in data:
        lines.extend(["## 📝 Questões", ""])
        for i, c in enumerate(data["correcoes"], 1):
            q_num = c.get("questao_numero", i)
            lines.append(f"### Questão {q_num}")
            lines.append(f"**Nota:** {c.get('nota', 0):.1f} / {c.get('nota_maxima', 1):.1f}")
            if c.get("feedback"):
                lines.append(f"> {c.get('feedback')}")
            lines.append("")
    
    if data.get("feedback"):
        lines.extend(["## 💬 Feedback", "", data["feedback"], ""])
    
    return lines


def _build_relatorio_md(data: Dict) -> List[str]:
    """Constrói Markdown para relatório final"""
    lines = []
    
    if data.get("aluno_nome"):
        lines.append(f"**Aluno:** {data.get('aluno_nome')}")
    if data.get("atividade_nome"):
        lines.append(f"**Atividade:** {data.get('atividade_nome')}")
    lines.append("")
    
    nota = data.get("nota_final", data.get("nota", 0))
    nota_max = data.get("nota_maxima", 10)
    percentual = data.get("percentual", (nota / nota_max * 100) if nota_max > 0 else 0)
    
    lines.extend([
        "## 📊 Resultado Final",
        "",
        f"| Métrica | Valor |",
        f"|---------|-------|",
        f"| Nota Final | {nota:.1f} / {nota_max:.1f} |",
        f"| Aproveitamento | {percentual:.1f}% |",
    ])
    
    if data.get("total_questoes"):
        lines.append(f"| Total de Questões | {data['total_questoes']} |")
        lines.append(f"| Corretas | {data.get('questoes_corretas', 0)} |")
    
    lines.append("")
    
    if data.get("recomendacoes"):
        lines.extend(["## 📌 Recomendações", ""])
        for rec in data["recomendacoes"]:
            lines.append(f"- {rec}")
        lines.append("")
    
    return lines


def _build_analise_md(data: Dict) -> List[str]:
    """Constrói Markdown para análise de habilidades"""
    lines = []
    
    habilidades = data.get("habilidades", {})
    
    if isinstance(habilidades, dict):
        dominadas = habilidades.get("dominadas", [])
        em_dev = habilidades.get("em_desenvolvimento", [])
        nao_dem = habilidades.get("nao_demonstradas", [])
    else:
        dominadas = data.get("habilidades_demonstradas", [])
        em_dev = []
        nao_dem = data.get("habilidades_faltantes", [])
    
    if dominadas:
        lines.extend(["## ✅ Habilidades Demonstradas", ""])
        for h in dominadas:
            nome = h.get("nome", h) if isinstance(h, dict) else h
            lines.append(f"- {nome}")
        lines.append("")
    
    if em_dev:
        lines.extend(["## 🔄 Em Desenvolvimento", ""])
        for h in em_dev:
            nome = h.get("nome", h) if isinstance(h, dict) else h
            lines.append(f"- {nome}")
        lines.append("")
    
    if nao_dem:
        lines.extend(["## ⚠️ Precisam de Atenção", ""])
        for h in nao_dem:
            nome = h.get("nome", h) if isinstance(h, dict) else h
            lines.append(f"- {nome}")
        lines.append("")
    
    return lines


def _build_generic_md(data: Dict) -> List[str]:
    """Constrói Markdown genérico"""
    lines = []
    
    for key, value in data.items():
        if key.startswith("_"):
            continue
        
        label = key.replace("_", " ").title()
        
        if isinstance(value, list):
            lines.extend([f"## {label}", ""])
            for item in value:
                if isinstance(item, dict):
                    lines.append(f"- {json.dumps(item, ensure_ascii=False)}")
                else:
                    lines.append(f"- {item}")
            lines.append("")
        elif isinstance(value, dict):
            lines.extend([f"## {label}", "", "```json", json.dumps(value, indent=2, ensure_ascii=False), "```", ""])
        else:
            lines.append(f"**{label}:** {value}")
    
    return lines


# ============================================================
# DOCX GENERATION (Basic)
# ============================================================

def generate_docx(data: Dict[str, Any], title: str = "Documento",
                  doc_type: str = "generic") -> bytes:
    """
    Gera DOCX a partir de dados estruturados.
    
    Args:
        data: Dicionário com dados
        title: Título do documento
        doc_type: Tipo de documento
    
    Returns:
        bytes: Conteúdo do DOCX
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Fallback: retornar como texto
        return generate_markdown(data, title, doc_type).encode('utf-8')
    
    doc = Document()
    
    # Título
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Gerar conteúdo baseado no Markdown (simples)
    md_content = generate_markdown(data, "", doc_type)
    
    for line in md_content.split('\n'):
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and ':**' in line:
            doc.add_paragraph(line.replace('**', ''))
        elif line.startswith('|'):
            # Skip table rows (simple)
            continue
        elif line.strip() and not line.startswith('#') and not line.startswith('---'):
            doc.add_paragraph(line)
    
    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    footer.runs[0].font.size = Pt(8)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ============================================================
# MAIN GENERATOR INTERFACE
# ============================================================

def generate_document(data: Dict[str, Any], 
                      format: OutputFormat,
                      title: str = "Documento",
                      doc_type: str = "generic") -> Union[bytes, str]:
    """
    Interface principal para gerar documento em qualquer formato.
    
    Args:
        data: Dados do documento
        format: Formato desejado (OutputFormat enum)
        title: Título do documento
        doc_type: Tipo de documento (correcao, relatorio_final, etc.)
    
    Returns:
        bytes ou str dependendo do formato
    """
    if format == OutputFormat.PDF:
        return generate_pdf(data, title, doc_type)
    elif format == OutputFormat.CSV:
        return generate_csv(data, doc_type=doc_type)
    elif format == OutputFormat.DOCX:
        return generate_docx(data, title, doc_type)
    elif format == OutputFormat.MD:
        return generate_markdown(data, title, doc_type)
    elif format == OutputFormat.JSON:
        return json.dumps(data, indent=2, ensure_ascii=False)
    else:
        raise ValueError(f"Formato não suportado: {format}")


def get_file_extension(format: OutputFormat) -> str:
    """Retorna a extensão de arquivo para um formato"""
    extensions = {
        OutputFormat.PDF: ".pdf",
        OutputFormat.CSV: ".csv",
        OutputFormat.DOCX: ".docx",
        OutputFormat.MD: ".md",
        OutputFormat.JSON: ".json",
    }
    return extensions.get(format, ".txt")
